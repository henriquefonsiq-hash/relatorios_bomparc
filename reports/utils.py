"""
Funções utilitárias compartilhadas entre os módulos de relatório.
Otimizado para Microsoft Word (Local) e LibreOffice (Streamlit Cloud).
"""
import os
import platform
import re
import subprocess
import tempfile

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from PIL import Image, ImageChops, ImageOps

if platform.system() == "Windows":
    import pythoncom
    import win32com.client

RED_VALUE = RGBColor(255, 0, 0)
SIGNATURE_EXTENSIONS = (".jpg", ".jpeg", ".png")
SIGNATURE_AREA_FILL = 0.82
SIGNATURE_CANVAS_HEIGHT_PX = 700


# ---------------------------------------------------------------------------
# Texto e Arquivos
# ---------------------------------------------------------------------------

def word_value(value):
    return str(value).upper()


def safe_filename_part(value):
    cleaned = re.sub(r'[<>:"/\\|?*]+', "-", str(value)).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip(". ") or "relatorio"


def parse_certificate_items(raw_items):
    items = []
    seen = set()
    for raw_item in re.split(r"[,;\n]+", raw_items):
        item_value = re.sub(r"(?i)^item\s+", "", raw_item.strip()).strip("- ")
        if not item_value or item_value in seen:
            continue
        items.append(item_value)
        seen.add(item_value)
    return items


# ---------------------------------------------------------------------------
# Assinaturas e Imagens
# ---------------------------------------------------------------------------

def list_signature_files(signature_dir):
    if not os.path.isdir(signature_dir):
        return []
    return sorted(
        os.path.join(signature_dir, name)
        for name in os.listdir(signature_dir)
        if name.lower().endswith(SIGNATURE_EXTENSIONS)
    )


def signature_label(path):
    if path is None:
        return "Selecione"
    return os.path.splitext(os.path.basename(path))[0]


def crop_signature_whitespace(image):
    image = ImageOps.exif_transpose(image).convert("RGBA")
    white = Image.new("RGBA", image.size, (255, 255, 255, 255))
    flattened = Image.alpha_composite(white, image).convert("RGB")
    diff = ImageChops.difference(
        flattened, Image.new("RGB", flattened.size, "white")
    ).convert("L")
    mask = diff.point(lambda pixel: 255 if pixel > 18 else 0)
    bbox = mask.getbbox()
    if bbox is None:
        return image
    left, top, right, bottom = bbox
    padding = 6
    return image.crop((
        max(left - padding, 0),
        max(top - padding, 0),
        min(right + padding, image.width),
        min(bottom + padding, image.height),
    ))


# ---------------------------------------------------------------------------
# Substituição Cirúrgica (Preserva Formatação do Template)
# ---------------------------------------------------------------------------

def replace_placeholder(paragraph, placeholder, value, color=None):
    """
    Substitui o placeholder no parágrafo de forma cirúrgica, 
    preservando a formatação individual de cada 'run' (pedaço de texto).
    Força o texto substituído a não ser negrito para manter elegância.
    """
    import re
    
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder not in full_text:
        return False
        
    val_str = word_value(value)
    start_indices = [m.start() for m in re.finditer(re.escape(placeholder), full_text)]
    
    for start_idx in reversed(start_indices):
        end_idx = start_idx + len(placeholder)
        current_pos = 0
        runs_involved = []
        for run in paragraph.runs:
            run_len = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_len
            if run_start < end_idx and run_end > start_idx:
                runs_involved.append({
                    "run": run,
                    "start": run_start,
                    "end": run_end,
                    "offset_start": max(0, start_idx - run_start),
                    "offset_end": min(run_len, end_idx - run_start)
                })
            current_pos = run_end
            
        if not runs_involved: continue
            
        first = runs_involved[0]
        r_first = first["run"]
        prefix = r_first.text[:first["offset_start"]]
        
        if len(runs_involved) == 1:
            suffix = r_first.text[first["offset_end"]:]
            r_first.text = prefix + val_str + suffix
            if color: r_first.font.color.rgb = color
            else: r_first.font.color.rgb = None
            r_first.font.bold = False  # Força fonte fina (estilo Ideal)
        else:
            last = runs_involved[-1]
            r_last = last["run"]
            suffix = r_last.text[last["offset_end"]:]
            r_first.text = prefix + val_str
            if color: r_first.font.color.rgb = color
            else: r_first.font.color.rgb = None
            r_first.font.bold = False  # Força fonte fina (estilo Ideal)
            for i in range(1, len(runs_involved) - 1):
                runs_involved[i]["run"].text = ""
            r_last.text = suffix

    return True


def process_replacements(doc, replacements):
    import re
    # Tags em vermelho (Normalizadas)
    RED_TAGS_LOWER = {
        "série", "data", "data de validade", 
        "ns (número de série)", "data de inspeção"
    }
    
    repl_lower = {k.lower().strip(): v for k, v in replacements.items()}
    tag_regex = re.compile(r'["“”]([^"“”]+)["“”]')

    def apply_to_paragraph(p):
        text = p.text
        if not text: return
        for match in tag_regex.finditer(text):
            full_tag = match.group(0)
            content = match.group(1).lower().strip()
            if content in repl_lower:
                value = repl_lower[content]
                color = RED_VALUE if content in RED_TAGS_LOWER else None
                replace_placeholder(p, full_tag, value, color)

    for p in doc.paragraphs: apply_to_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs: apply_to_paragraph(p)


def set_cell_text(cell, value, color=None):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.text = word_value(value)
    if color: paragraph.runs[0].font.color.rgb = color


# ---------------------------------------------------------------------------
# Conversão PDF
# ---------------------------------------------------------------------------

def convert_to_pdf(docx_path, pdf_path):
    if platform.system() == "Windows":
        pythoncom.CoInitialize()
        word = None; doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        finally:
            if doc is not None: doc.Close(False)
            if word is not None: word.Quit()
            pythoncom.CoUninitialize()
    else:
        # Comando para Streamlit Cloud (Linux/LibreOffice)
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", 
            os.path.abspath(docx_path), 
            "--outdir", os.path.dirname(os.path.abspath(pdf_path))
        ], check=True)

    if not os.path.exists(pdf_path):
        raise RuntimeError("o arquivo PDF não foi criado.")


def format_conversion_error(error):
    if platform.system() == "Windows":
        return (
            "Não foi possível converter para PDF localmente. "
            "Feche o Word e tente novamente. "
            f"Erro: {error}"
        )
    return f"Erro na conversão online (Streamlit): {error}"


# ---------------------------------------------------------------------------
# Manipulação de Células e Documentos
# ---------------------------------------------------------------------------

def add_signature_to_cell(cell, signature_path):
    if not signature_path: return
    with Image.open(signature_path) as source:
        signature = crop_signature_whitespace(source)
    max_w, max_h = 1.55, 0.58
    scale = min(max_w / signature.width, max_h / signature.height)
    fd, fitted_path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        signature.save(fitted_path)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs: run.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(fitted_path, width=Inches(signature.width * scale), height=Inches(signature.height * scale))
    finally:
        if os.path.exists(fitted_path): os.remove(fitted_path)


def add_photo_to_cell(cell, image_file):
    if image_file is None: return
    image_file.seek(0)
    with Image.open(image_file) as source:
        image = ImageOps.exif_transpose(source).convert("RGB")
    max_w, max_h = 6.25, 5.25
    scale = min(max_w / image.width, max_h / image.height)
    fd, image_path = tempfile.mkstemp(suffix=".jpg")
    os.close(fd)
    try:
        image.save(image_path, quality=95)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs: run.text = ""
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(image_path, width=Inches(image.width * scale), height=Inches(image.height * scale))
    finally:
        if os.path.exists(image_path): os.remove(image_path)


def append_document(master, sub):
    master.add_page_break()
    for element in sub.element.body:
        if element.tag.endswith('sectPr'): continue
        master.element.body.append(element)
