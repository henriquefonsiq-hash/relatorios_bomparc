"""
FOR-BPC-5137 — Relatório de Inspeção Visual de Equipamento

Layout: formulário à esquerda | pré-visualização ao vivo à direita.
A prévia atualiza automaticamente a cada alteração no formulário,
usando mammoth (docx→HTML, ~100ms) sem necessidade de botão.

Roteamento de máscaras:
  APROVADO  → template.docx             (FOR-BPC-5137)
  REPROVADO → template_reprovados.docx  (FOR-BPC-5112)
  END LP    → template_5113_lp.docx     (FOR-BPC-5113) - Anexado como página extra
  END PM    → template_5114_pm.docx     (FOR-BPC-5114) - Anexado como página extra
"""

import base64
import io
import os
import re
import tempfile
import zipfile
from datetime import date, timedelta

import docx
import mammoth
import streamlit as st
import streamlit.components.v1 as components
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from PIL import Image, ImageOps

from client_registry import load_client_registry
from equipment_registry import load_equipment_registry
from reports.utils import (
    RED_VALUE,
    SIGNATURE_AREA_FILL,
    SIGNATURE_CANVAS_HEIGHT_PX,
    add_photo_to_cell,
    add_signature_to_cell,
    convert_to_pdf,
    crop_signature_whitespace,
    format_conversion_error,
    list_signature_files,
    parse_certificate_items,
    replace_placeholder,
    process_replacements,
    safe_filename_part,
    set_cell_text,
    signature_label,
    word_value,
    append_document,
)
from reports.engine_5112 import build_5112
from reports.engine_5113 import build_5113
from reports.engine_5114 import build_5114

# ── Constantes ───────────────────────────────────────────────────────────────

APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TECHNICIAN_SIGNATURE_DIR    = os.path.join(APP_DIR, "assinatura técnicos")
QUALITY_SIGNATURE_DIR       = os.path.join(APP_DIR, "assinatura qualidade")
TEMPLATE_APROVADO           = os.path.join(APP_DIR, "template.docx")
TEMPLATE_REPROVADO          = os.path.join(APP_DIR, "template_reprovados.docx")
QUALITY_SIGNATURE_TARGET    = "media/image1.jpeg"
TECHNICIAN_SIGNATURE_TARGET = "media/image3.jpeg"

IDEAL_TABLE_GRID_WIDTHS            = [7620, 1688465, 469265, 407035, 747395, 361315, 739140, 786130, 1710055]
IDEAL_EQUIPMENT_DETAIL_CELL_WIDTHS = [3398, 3551, 3931]
IDEAL_SIGNATURE_CELL_WIDTHS        = [2671, 1380, 2910, 1238, 2693]

# ── Utilitários — template APROVADO (5137) ───────────────────────────────────

def paragraph_has_image_target(paragraph, target_ref):
    for blip in paragraph._element.iter(qn("a:blip")):
        rel_id = blip.get(qn("r:embed"))
        if rel_id and paragraph.part.rels[rel_id].target_ref == target_ref:
            return True
    return False

def first_descendant(element, tag):
    return next(element.iter(qn(tag)), None)

def make_signature_fit_canvas(signature_path, target_cx, target_cy):
    if not target_cx or not target_cy:
        return signature_path
    with Image.open(signature_path) as source:
        signature = crop_signature_whitespace(source)
    aspect_ratio = target_cx / target_cy
    canvas_height = SIGNATURE_CANVAS_HEIGHT_PX
    canvas_width  = max(1, int(round(canvas_height * aspect_ratio)))
    max_width  = max(1, int(canvas_width  * SIGNATURE_AREA_FILL))
    max_height = max(1, int(canvas_height * SIGNATURE_AREA_FILL))
    scale    = min(max_width / signature.width, max_height / signature.height)
    new_size = (max(1, int(round(signature.width * scale))), max(1, int(round(signature.height * scale))))
    signature = signature.resize(new_size, Image.Resampling.LANCZOS)
    canvas = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
    canvas.alpha_composite(signature, (int((canvas_width - signature.width) / 2), int((canvas_height - signature.height) / 2)))
    fd, fitted_path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    canvas.save(fitted_path)
    return fitted_path

def keep_cell_text_on_one_line(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")

def apply_ideal_table_dimensions(doc):
    if not doc.tables:
        return
    table = doc.tables[0]
    for col, width in zip(table._tbl.tblGrid.gridCol_lst, IDEAL_TABLE_GRID_WIDTHS):
        col.w = width
    if len(table.rows) <= 17:
        return
    from docx.table import _Cell as DocxCell
    for row_idx in (6, 7):
        for tc, w in zip(list(table.rows[row_idx]._tr.tc_lst), IDEAL_EQUIPMENT_DETAIL_CELL_WIDTHS):
            set_cell_width(DocxCell(tc, table), w)
    for row_idx in (16, 17):
        for tc, w in zip(list(table.rows[row_idx]._tr.tc_lst), IDEAL_SIGNATURE_CELL_WIDTHS):
            set_cell_width(DocxCell(tc, table), w)

def replace_signature_aprovado(doc, signature_path, target_ref):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph_has_image_target(paragraph, target_ref):
                        continue
                    for drawing in paragraph._element.iter(qn("w:drawing")):
                        for blip in drawing.iter(qn("a:blip")):
                            rel_id = blip.get(qn("r:embed"))
                            if not rel_id or paragraph.part.rels[rel_id].target_ref != target_ref:
                                continue
                            extent    = first_descendant(drawing, "wp:extent")
                            target_cx = int(extent.get("cx")) if extent is not None else None
                            target_cy = int(extent.get("cy")) if extent is not None else None
                            fitted_path = make_signature_fit_canvas(signature_path, target_cx, target_cy)
                            try:
                                new_rel_id, _ = paragraph.part.get_or_add_image(fitted_path)
                            finally:
                                if fitted_path != signature_path and os.path.exists(fitted_path):
                                    os.remove(fitted_path)
                            blip.set(qn("r:embed"), new_rel_id)
                            anchor = first_descendant(drawing, "wp:anchor")
                            if anchor is not None:
                                anchor.set("behindDoc", "1")
                            return True
    return False

def validate_template_aprovado(doc):
    if not doc.tables:
        return "O template.docx não possui tabela principal."
    table = doc.tables[0]
    if len(table.rows) <= 17:
        return "A tabela principal do template.docx precisa ter pelo menos 18 linhas."
    for row_idx in (16, 17):
        if len(table.rows[row_idx].cells) <= 7:
            return "As linhas de assinatura precisam ter pelo menos 8 colunas."
    image_targets = {r.target_ref for r in doc.part.rels.values() if r.target_ref.startswith("media/image")}
    missing = [t for t in (QUALITY_SIGNATURE_TARGET, TECHNICIAN_SIGNATURE_TARGET) if t not in image_targets]
    if missing:
        return "O template.docx não possui as imagens-base de assinatura: " + ", ".join(missing)
    return None


# ── Build dos documentos (retornam Objeto Document) ──────────────────────────

def _build_aprovado_doc(campos):
    if not os.path.exists(TEMPLATE_APROVADO):
        raise FileNotFoundError(f"'template.docx' não encontrado em {APP_DIR}.")

    item_atual    = campos["itens_certificado"][0]
    certif        = f"{campos['numero_certificado']}-{item_atual}"
    os_digits     = campos['numero_certificado'][:4]

    replacements = {
        # Tags Simples
        "Relatório":             certif,
        "Número do Certificado": certif,
        "Item(s)":               item_atual,
        "Cliente":               campos["cliente"],
        "Embarcação":            campos["embarcacao"],
        "Endereço":              campos["endereco"],
        "Equipamento":           campos["equipamento_str"],
        "Série":                 campos["ns"],
        "Data":                  campos["data_str"],
        "Data de Validade":      campos["data_validade_str"],
        "Critério de Aceitação": campos["criterio"],
        "Carga de Trabalho":     campos["carga_trabalho"],
        "Capacidade":            f"{campos['capac']:g}",
        "Unidade":               campos["unidade_capac"],
        "Dimensão":              campos["dimensao"],
        "Quantidade":            f"{campos['quantidade']:02d} UNIDADE",
        "Matéria Prima":         campos["materia_prima"],
        "Teste":                 campos["teste"],
        "END":                   campos["end"],
        "Laudo Final":           campos["aprov"],
        "Observações":           campos["obs"],
        "Ordem de Serviço":      os_digits,

        # Sinônimos Longos (para bater com os nomes dos campos na tela)
        "NS (Número de Série)":             campos["ns"],
        "Data de Inspeção":                 campos["data_str"],
        "Data de Inspeção 2":               campos["data_str"],
        "Critério de Aceitação / Norma":    campos["criterio"],
        "Data de Validade (Validity Date)": campos["data_validade_str"],
    }

    doc_out = docx.Document(TEMPLATE_APROVADO)
    err = validate_template_aprovado(doc_out)
    if err:
        raise ValueError(err)

    if campos.get("assinatura_qualidade"):
        replace_signature_aprovado(doc_out, campos["assinatura_qualidade"], QUALITY_SIGNATURE_TARGET)
    if campos.get("assinatura_tecnico"):
        replace_signature_aprovado(doc_out, campos["assinatura_tecnico"], TECHNICIAN_SIGNATURE_TARGET)

    process_replacements(doc_out, replacements)

    sig_table = doc_out.tables[0]
    for cell_idx in (2, 7):
        keep_cell_text_on_one_line(sig_table.rows[16].cells[cell_idx])
        keep_cell_text_on_one_line(sig_table.rows[17].cells[cell_idx])
    apply_ideal_table_dimensions(doc_out)

    return doc_out


def _build_documento(campos):
    """Orquestra a criação do documento completo, incluindo páginas extras de END."""
    # 1. Gera o documento base (Visual ou Reprovado)
    if campos["aprov"] == "REPROVADO":
        main_doc = build_5112(campos)
    else:
        main_doc = _build_aprovado_doc(campos)
    
    # 2. Anexa páginas adicionais para cada END selecionado
    if "LP" in campos["end_selected"]:
        doc_lp = build_5113(campos)
        append_document(main_doc, doc_lp)
        
    if "PM" in campos["end_selected"]:
        doc_pm = build_5114(campos)
        append_document(main_doc, doc_pm)

    # 3. Salva em bytes para prévia ou download
    buf = io.BytesIO()
    main_doc.save(buf)
    return buf.getvalue()


# ── Prévia ao vivo: docx → HTML via mammoth ──────────────────────────────────

_PREVIEW_CSS = """
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    font-family: Arial, sans-serif;
    font-size: 10px;
    background: #f0f0f0;
    padding: 16px;
  }
  .page {
    background: white;
    width: 100%;
    padding: 18px 20px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.15);
    border-radius: 3px;
    margin-bottom: 20px;
  }
  p  { margin: 2px 0; line-height: 1.4; }
  strong { font-weight: bold; }
  em { font-style: italic; }
  table {
    border-collapse: collapse;
    width: 100%;
    margin: 6px 0;
    font-size: 9px;
  }
  td, th {
    border: 1px solid #333;
    padding: 3px 5px;
    vertical-align: top;
  }
  img {
    max-width: 100%;
    height: auto;
  }
  /* destaca campos preenchidos em vermelho (NS, datas) */
  span.filled-red { color: #c00; font-weight: bold; }
</style>
"""

def _docx_bytes_to_html(docx_bytes):
    """Converte bytes de docx para HTML usando mammoth."""
    def convert_image(image):
        with image.open() as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        ct  = image.content_type or "image/png"
        return {"src": f"data:{ct};base64,{b64}"}

    result = mammoth.convert_to_html(
        io.BytesIO(docx_bytes),
        convert_image=mammoth.images.img_element(convert_image),
    )
    return result.value


def _render_preview(campos):
    """Gera e exibe a prévia HTML do documento no painel direito."""
    try:
        docx_bytes = _build_documento(campos)
        html_body  = _docx_bytes_to_html(docx_bytes)
        full_html  = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">{_PREVIEW_CSS}</head>
<body><div class="page">{html_body}</div></body></html>"""
        components.html(full_html, height=780, scrolling=True)
    except Exception as e:
        st.warning(f"Pré-visualização indisponível: {e}")


# ── Monta campos sem validação (para prévia parcial) ─────────────────────────

def _campos_parciais(
    numero_certificado, item_certificado, cliente, embarcacao, endereco,
    data_inspecao, ns, item, criterio, capac, unidade_capac, dimensao,
    quantidade, materia_prima, teste, end, end_selected, aprov, obs,
    assinatura_qualidade, signature_tecnico,
    descricao_insuficiencia, foto_reprovado,
):
    """Monta o dicionário de campos usando defaults seguros para valores vazios."""

    num   = numero_certificado.strip().strip("-") or "XXXXXXXX"
    itens = parse_certificate_items(item_certificado) or ["00"]
    uni   = unidade_capac.strip().upper() or "TON"
    cap   = capac if capac is not None else 0.0
    qty   = quantidade if quantidade is not None else 1

    data_str     = data_inspecao.strftime("%d/%m/%Y") if data_inspecao else "__/__/____"
    data_arq_str = data_inspecao.strftime("%d-%m-%Y")  if data_inspecao else "00-00-0000"
    
    # Lógica de Validade: Inspeção + 1 ano (com tratamento para ano bissexto)
    if data_inspecao:
        try:
            validade_dt = data_inspecao.replace(year=data_inspecao.year + 1)
        except ValueError:
            validade_dt = data_inspecao.replace(year=data_inspecao.year + 1, day=28)
        data_val_str = validade_dt.strftime("%d/%m/%Y")
    else:
        data_val_str = "__/__/____"

    if uni == "TON" and cap > 0:
        carga = f"{int(cap * 1000):,} KG".replace(",", ".")
    elif cap > 0:
        carga = f"{cap:g} {uni}"
    else:
        carga = ""

    equip_str = f"{item.upper()} {cap:g} {uni} X {dimensao.upper()}" if item and cap else item or ""

    return {
        "numero_certificado":    num,
        "itens_certificado":     itens,
        "cliente":               cliente or "",
        "embarcacao":            embarcacao or "",
        "endereco":              endereco or "",
        "ns":                    ns or "",
        "data_str":              data_str,
        "data_arquivo_str":      data_arq_str,
        "data_validade_str":     data_val_str,
        "item":                  item or "",
        "criterio":              criterio or "",
        "capac":                 cap,
        "unidade_capac":         uni,
        "dimensao":              dimensao or "",
        "quantidade":            qty,
        "materia_prima":         materia_prima or "",
        "teste":                 teste or "",
        "carga_trabalho":        carga,
        "equipamento_str":       equip_str,
        "end":                   end or "",
        "end_selected":          end_selected or [],
        "aprov":                 aprov if aprov else "APROVADO",
        "obs":                   obs or "",
        "assinatura_qualidade":  assinatura_qualidade,
        "assinatura_tecnico":    signature_tecnico,
        "descricao_insuficiencia": descricao_insuficiencia or "",
        "foto_reprovado":        foto_reprovado,
    }


# ── Validação estrita (para geração final) ───────────────────────────────────

def _validar_campos(campos_p):
    erros = []
    if campos_p["numero_certificado"] == "XXXXXXXX":
        erros.append("Preencha o Número do certificado.")
    if campos_p["itens_certificado"] == ["00"]:
        erros.append("Preencha pelo menos um Item.")
    if not campos_p["unidade_capac"]:
        erros.append("Preencha a Unidade da capacidade.")
    if not campos_p["criterio"]:
        erros.append("Preencha o Critério de Aceitação.")
    if campos_p["data_str"] == "__/__/____":
        erros.append("Preencha a Data da Inspeção.")
    if not campos_p["item"]:
        erros.append("Preencha o Equipamento.")
    if campos_p["capac"] == 0.0:
        erros.append("Preencha a Capacidade.")
    return erros


# ── Geração do ZIP final ──────────────────────────────────────────────────────

def _gerar_zip(campos):
    with tempfile.TemporaryDirectory() as temp_dir:
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for item_atual in campos["itens_certificado"]:
                campos_item = {**campos, "itens_certificado": [item_atual]}
                docx_bytes = _build_documento(campos_item)
                sufixo = "REPROVADO" if campos["aprov"] == "REPROVADO" else ""
                certif_str = f"{campos['numero_certificado']}-{item_atual}" + ("-RI" if sufixo else "")
                if "LP" in campos["end_selected"]: certif_str += "-LP"
                if "PM" in campos["end_selected"]: certif_str += "-PM"
                nome_base  = " ".join(safe_filename_part(p) for p in filter(None, (certif_str, sufixo, campos["item"], campos["data_arquivo_str"])))
                docx_path = os.path.join(temp_dir, f"{nome_base}.docx")
                pdf_path  = os.path.join(temp_dir, f"{nome_base}.pdf")
                with open(docx_path, "wb") as f: f.write(docx_bytes)
                convert_to_pdf(docx_path, pdf_path)
                zip_file.write(docx_path, os.path.basename(docx_path))
                zip_file.write(pdf_path,  os.path.basename(pdf_path))
        zip_buffer.seek(0)
        return zip_buffer.getvalue()


# ════════════════════════════════════════════════════════════════════════════
# INTERFACE
# ════════════════════════════════════════════════════════════════════════════

st.set_page_config(layout="wide")
st.title("Gerador Automático de Relatório — BOMPARC")

equipment_registry = load_equipment_registry()
equipment_options  = sorted(equipment_registry)
client_registry    = load_client_registry()
client_options     = sorted(client_registry)

col_form, col_preview = st.columns([1, 1], gap="large")

with col_form:
    st.subheader("1. Identificação e Local")
    c1, c2, c3 = st.columns(3)
    cliente = c1.selectbox("Cliente", ["Selecione", *client_options, "Outro"], key="cliente_item")
    if cliente == "Selecione": cliente = ""
    if cliente == "Outro": cliente = c1.text_input("Especifique o Cliente", "", key="cliente_outro")
    vessel_options = sorted(client_registry.get(cliente, {}))
    embarcacao = c2.selectbox("Embarcação", ["Selecione", *vessel_options, "Outro"], key="embarcacao_item")
    if embarcacao == "Selecione": embarcacao = ""
    if embarcacao == "Outro": embarcacao = c2.text_input("Especifique a Embarcação", "", key="embarcacao_outro")
    if st.session_state.get("_ult_emb_cli") != (cliente, embarcacao):
        st.session_state["endereco_cliente"] = client_registry.get(cliente, {}).get(embarcacao, "")
        st.session_state["_ult_emb_cli"] = (cliente, embarcacao)
    endereco = c3.text_input("Endereço", key="endereco_cliente")
    c4, c5 = st.columns(2)
    numero_certificado = c4.text_input("Número do certificado", "")
    item_certificado   = c5.text_input("Item(s)", "")

    st.subheader("2. Dados da Inspeção e Equipamento")
    c6, c7, c8 = st.columns(3)
    data_inspecao = c6.date_input("Data da Inspeção", value=None, format="DD/MM/YYYY")
    ns            = c7.text_input("NS (Número de Série)", "")
    item = c8.selectbox("Equipamento", ["Selecione", *equipment_options, "Outro"], key="equip_item")
    if item == "Selecione": item = ""
    if item == "Outro": item = st.text_input("Especifique o Equipamento", "", key="equip_outro")
    if st.session_state.get("_ult_equip") != item:
        cfg = equipment_registry.get(item, {})
        st.session_state["criterio_aceitacao"] = cfg.get("criterio_aceitacao", "")
        st.session_state["materia_prima"]       = cfg.get("materia_prima", "")
        st.session_state["_ult_equip"] = item
    criterio = st.text_input("Critério de Aceitação / Norma", key="criterio_aceitacao")
    c9, c10, c11, c12 = st.columns(4)
    capac         = c9.number_input("Capacidade", min_value=0.1, value=None, step=0.1)
    unidade_capac = c10.text_input("Unidade", "")
    dimensao      = c11.text_input("Dimensão", "")
    quantidade    = c12.number_input("Quantidade", min_value=1, value=None, step=1)
    c13, c14 = st.columns(2)
    materia_prima = c13.text_input("Matéria Prima", key="materia_prima")
    teste         = c14.text_input("Teste", "")

    st.subheader("3. Parâmetros e Parecer")
    c15, c16 = st.columns(2)
    end1 = c15.selectbox("Página Extra 1 (END)", ["N/A", "LP", "PM"])
    end_selected = []
    if end1 != "N/A":
        end_selected.append(end1)
        restante = "PM" if end1 == "LP" else "LP"
        end2 = c16.selectbox("Página Extra 2 (END)", ["N/A", restante])
        if end2 != "N/A": end_selected.append(end2)
    end = ", ".join(end_selected) if end_selected else "N/A"
    c15_2, c16_2 = st.columns(2)
    aprov = c15_2.selectbox("Laudo Final (APROV)", ["Selecione", "APROVADO", "REPROVADO"])
    if aprov == "Selecione": aprov = ""
    obs   = c16_2.text_area("Observações (OBS)", "")

    descricao_insuficiencia = ""
    foto_reprovado = []
    if aprov == "REPROVADO":
        st.divider()
        st.subheader("4. Dados de Reprovados (FOR-BPC-5112)")
        descricao_insuficiencia = st.text_area("Descrição da Insuficiência", "")
        foto_reprovado = st.file_uploader("Fotos do equipamento reprovado", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

    st.divider()
    quality_files = list_signature_files(QUALITY_SIGNATURE_DIR)
    assinatura_qualidade = st.selectbox("Assinatura — Controle de Qualidade", [None, *quality_files], format_func=signature_label) if quality_files else None
    tech_files = list_signature_files(TECHNICIAN_SIGNATURE_DIR)
    assinatura_tecnico = st.selectbox("Assinatura — Técnico de Inspeção", [None, *tech_files], format_func=signature_label) if tech_files else None

    st.divider()
    submit_button = st.button("📄 Gerar Relatório Unificado (Word + PDF)", type="primary", use_container_width=True)

campos = _campos_parciais(numero_certificado, item_certificado, cliente, embarcacao, endereco, data_inspecao, ns, item, criterio, capac, unidade_capac, dimensao, quantidade, materia_prima, teste, end, end_selected, aprov, obs, assinatura_qualidade, assinatura_tecnico, descricao_insuficiencia, foto_reprovado)

with col_preview:
    mascara_label = "FOR-BPC-5112" if campos["aprov"] == "REPROVADO" else "FOR-BPC-5137"
    if campos["end_selected"]: mascara_label += " + " + " & ".join(campos["end_selected"])
    st.subheader(f"Pré-visualização · {mascara_label}")
    _render_preview(campos)

if submit_button:
    erros = _validar_campos(campos)
    if erros:
        with col_form:
            for e in erros: st.error(e)
    else:
        with col_form:
            with st.spinner("Gerando documento unificado..."):
                try:
                    zip_bytes = _gerar_zip(campos)
                    st.success("Relatório generado com sucesso!")
                    st.download_button(label="⬇️ Baixar Word e PDF (.zip)", data=zip_bytes, file_name=f"Relatorio_{campos['numero_certificado']}.zip", mime="application/zip")
                except Exception as e: st.error(format_conversion_error(e))
