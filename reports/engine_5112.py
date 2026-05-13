"""
Engine para a máscara FOR-BPC-5112 — Relatório de Equipamentos Reprovados.
Esta máscara é usada como fallback para diversos tipos de inspeção.
Inclui suporte a mosaico de fotos na página 2.
"""
import io
import os
import math
import docx
from PIL import Image
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reports.utils import (
    add_photo_to_cell,
    add_signature_to_cell,
    process_replacements,
)

def validate_template_5112(doc):
    if len(doc.tables) < 8:
        return "O template_reprovados.docx precisa possuir as tabelas esperadas da máscara FOR-BPC-5112."
    return None

def get_rows_count(count):
    if count <= 2: return 1
    if count <= 6: return 2
    if count <= 12: return 3
    if count <= 20: return 4
    return math.ceil(math.sqrt(count))

def apply_mosaic_to_cell(cell, photos):
    """Aplica o mosaico de fotos em uma célula de tabela do Word."""
    cell.text = ""
    for p in cell.paragraphs:
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0

    if not photos:
        cell.text = "Nenhuma foto anexada."
        return

    if not isinstance(photos, list):
        add_photo_to_cell(cell, photos)
        return
    
    count = len(photos)
    if count == 0:
        cell.text = "Nenhuma foto anexada."
        return
    if count == 1:
        add_photo_to_cell(cell, photos[0])
        return

    num_rows = get_rows_count(count)
    photo_data = []
    for f in photos:
        img = Image.open(f)
        photo_data.append({"file": f, "ratio": img.width / img.height})

    rows_data = []
    remaining = count
    photo_index = 0
    for i in range(num_rows):
        items_in_this_row = int(remaining / (num_rows - i) + 0.999)
        row_items = photo_data[photo_index:photo_index + items_in_this_row]
        rows_data.append(row_items)
        photo_index += items_in_this_row
        remaining -= items_in_this_row

    TOTAL_WIDTH_INCHES = 6.4 
    GAP_INCHES = 0.05 

    for row_items in rows_data:
        n_items = len(row_items)
        available_width = TOTAL_WIDTH_INCHES - (GAP_INCHES * (n_items - 1))
        sum_ratios = sum(item["ratio"] for item in row_items)
        
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0
        
        for idx, item in enumerate(row_items):
            item_width = available_width * (item["ratio"] / sum_ratios)
            run = p.add_run()
            run.add_picture(item["file"], width=Inches(item_width))
            if idx < n_items - 1:
                run_gap = p.add_run(" ")
                run_gap.font.size = docx.shared.Pt(2)

def build_5112(campos):
    app_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    template_path = os.path.join(app_dir, "template_reprovados.docx")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"'template_reprovados.docx' não encontrado em {app_dir}.")

    item_atual = campos["itens_certificado"][0]
    certif     = f"{campos['numero_certificado']}-{item_atual}-RI"

    replacements = {
        "Relatório":             certif,
        "Número do Certificado": certif,
        "Cliente":               campos["cliente"],
        "Embarcação":            campos["embarcacao"],
        "Endereço":              campos["endereco"],
        "Equipamento":           campos["item"],
        "Série":                 campos["ns"],
        "Data":                  campos["data_str"],
        "Data de Inspeção 2":    campos["data_str"],
        "Critério de Aceitação": campos["criterio"],
        "Carga de Trabalho":     campos["carga_trabalho"],
        "Capacidade":            f"{campos['capac']:g}",
        "Unidade":               campos["unidade_capac"],
        "Dimensão":              campos["dimensao"],
        "Quantidade":            f"{campos['quantidade']:02d} UNIDADE",
        "Matéria Prima":         campos["materia_prima"],
        "Descrição da Insuficiência": campos.get("descricao_insuficiencia", ""),
    }

    doc_out = docx.Document(template_path)
    err = validate_template_5112(doc_out)
    if err:
        raise ValueError(err)

    process_replacements(doc_out, replacements)
    apply_mosaic_to_cell(doc_out.tables[5].rows[0].cells[0], campos.get("foto_reprovado"))

    for table in doc_out.tables:
        for row in table.rows:
            for cell in row.cells:
                if "ASSINATURA" in cell.text.upper() or "SIGNATURE" in cell.text.upper():
                    if "QUALIDADE" in cell.text.upper():
                        add_signature_to_cell(cell, campos.get("assinatura_qualidade"))
                    else:
                        add_signature_to_cell(cell, campos.get("assinatura_tecnico"))

    return doc_out
