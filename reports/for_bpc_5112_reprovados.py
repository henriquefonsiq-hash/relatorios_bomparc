import io
import os
import platform
import re
import subprocess
import tempfile
import zipfile
from datetime import date

import docx
import streamlit as st
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageChops, ImageOps

from client_registry import load_client_registry
from equipment_registry import load_equipment_registry


if platform.system() == "Windows":
    import pythoncom
    import win32com.client


APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TECHNICIAN_SIGNATURE_DIR = os.path.join(APP_DIR, "assinatura técnicos")
QUALITY_SIGNATURE_DIR = os.path.join(APP_DIR, "assinatura qualidade")
SIGNATURE_EXTENSIONS = (".jpg", ".jpeg", ".png")
VALUE_FONT_SIZE = Pt(9)
RED_VALUE = RGBColor(255, 0, 0)


st.title("Gerador Automático de Relatório - BOMPARC")
st.markdown("Preencha o formulário abaixo para gerar o relatório de equipamentos reprovados (Word e PDF).")


def word_value(value):
    return str(value).upper()


def list_signature_files(signature_dir):
    if not os.path.isdir(signature_dir):
        return []
    return sorted(
        os.path.join(signature_dir, name)
        for name in os.listdir(signature_dir)
        if name.lower().endswith(SIGNATURE_EXTENSIONS)
    )


def signature_label(path):
    if path is None:
        return "Selecione"
    return os.path.splitext(os.path.basename(path))[0]


def safe_filename_part(value):
    cleaned = re.sub(r'[<>:"/\\\\|?*]+', "-", str(value)).strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip(". ") or "relatorio"


def parse_certificate_items(raw_items):
    items = []
    seen = set()
    for raw_item in re.split(r"[,;\n]+", raw_items):
        item_value = re.sub(r"(?i)^item\s+", "", raw_item.strip()).strip("- ")
        if not item_value or item_value in seen:
            continue
        items.append(item_value)
        seen.add(item_value)
    return items


def convert_to_pdf(docx_path, pdf_path):
    if platform.system() == "Windows":
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        finally:
            if doc is not None:
                doc.Close(False)
            if word is not None:
                word.Quit()
            pythoncom.CoUninitialize()
    else:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                os.path.abspath(docx_path),
                "--outdir",
                os.path.dirname(os.path.abspath(pdf_path)),
            ],
            check=True,
        )

    if not os.path.exists(pdf_path):
        raise RuntimeError("o arquivo PDF não foi criado.")


def format_conversion_error(error):
    if platform.system() == "Windows":
        return (
            "Não foi possível converter o relatório para PDF. "
            "Verifique se o Microsoft Word está instalado e fechado, e tente novamente. "
            f"Detalhe técnico: {error}"
        )
    return (
        "Não foi possível converter o relatório para PDF. "
        "Verifique se o LibreOffice está instalado no ambiente. "
        f"Detalhe técnico: {error}"
    )


def set_value_run_style(run, color=None):
    run.font.name = "Lato"
    run.font.size = VALUE_FONT_SIZE
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for font_attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(font_attr), "Lato")
    if color is not None:
        run.font.color.rgb = color
    else:
        color_element = r_pr.find(qn("w:color"))
        if color_element is not None:
            r_pr.remove(color_element)


def set_paragraph_text_preserving_first_run(paragraph, value, color=None):
    first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    first_run.text = word_value(value)
    set_value_run_style(first_run, color)
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_value_after_label(paragraph, label, value, color=None):
    full_text = paragraph.text
    if label not in full_text:
        return False

    value_start = full_text.index(label) + len(label)
    value_start += len(full_text[value_start:]) - len(full_text[value_start:].lstrip())
    value_end = len(full_text)

    if value_start >= value_end:
        value_run = paragraph.add_run(" " + word_value(value))
        set_value_run_style(value_run, color)
        return True

    cursor = 0
    inserted = False
    for run in paragraph.runs:
        original_text = run.text
        run_start = cursor
        run_end = cursor + len(original_text)
        cursor = run_end

        if run_end <= value_start:
            continue
        if run_start >= value_end:
            run.text = ""
            continue

        if not inserted:
            keep_before = original_text[: max(value_start - run_start, 0)]
            keep_after = original_text[max(value_end - run_start, 0) :]
            run.text = keep_before + word_value(value) + keep_after
            set_value_run_style(run, color)
            inserted = True
        else:
            keep_after = original_text[max(value_end - run_start, 0) :]
            run.text = keep_after

    return True


def replace_in_paragraphs(paragraphs, replacements, date_value):
    red_labels = {
        "SÉRIE (Serial Number):",
        "DATA DA INSPEÇÃO (Date):",
    }
    for paragraph in paragraphs:
        text = paragraph.text
        replaced_label = False
        for label, value in replacements.items():
            color = RED_VALUE if label in red_labels else None
            if replace_value_after_label(paragraph, label, value, color):
                replaced_label = True
                break

        if not replaced_label and (re.fullmatch(r"\s*/?\s*\d{4}\s*", text) or re.search(r"/\d{4}", text)):
            set_paragraph_text_preserving_first_run(paragraph, date_value, RED_VALUE)


def set_cell_text(cell, value):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    set_paragraph_text_preserving_first_run(paragraph, value)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text_preserving_first_run(paragraph, "")


def remove_empty_trailing_paragraphs(doc):
    body = doc._element.body

    for child in reversed(list(body)):
        if child.tag == qn("w:sectPr"):
            continue

        if child.tag != qn("w:p"):
            break

        has_text = any((text.text or "").strip() for text in child.iter(qn("w:t")))
        has_drawing = any(child.iter(qn("w:drawing")))
        has_page_break = any(
            break_element.get(qn("w:type")) == "page"
            for break_element in child.iter(qn("w:br"))
        )

        if has_text or has_drawing or has_page_break:
            break

        body.remove(child)


def crop_signature_whitespace(image):
    image = ImageOps.exif_transpose(image).convert("RGBA")
    white = Image.new("RGBA", image.size, (255, 255, 255, 255))
    flattened = Image.alpha_composite(white, image).convert("RGB")
    diff = ImageChops.difference(flattened, Image.new("RGB", flattened.size, "white")).convert("L")
    bbox = diff.point(lambda pixel: 255 if pixel > 18 else 0).getbbox()
    if bbox is None:
        return image

    left, top, right, bottom = bbox
    padding = 6
    return image.crop(
        (
            max(left - padding, 0),
            max(top - padding, 0),
            min(right + padding, image.width),
            min(bottom + padding, image.height),
        )
    )


def add_signature_to_cell(cell, signature_path):
    if not signature_path:
        return

    with Image.open(signature_path) as source:
        signature = crop_signature_whitespace(source)

    max_width_inches = 1.55
    max_height_inches = 0.58
    scale = min(max_width_inches / signature.width, max_height_inches / signature.height)
    width_inches = signature.width * scale
    height_inches = signature.height * scale

    fd, fitted_path = tempfile.mkstemp(suffix=".png")
    os.close(fd)
    try:
        signature.save(fitted_path)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""
            for drawing in list(paragraph._element.iter(qn("w:drawing"))):
                drawing.getparent().remove(drawing)

        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(
            fitted_path,
            width=Inches(width_inches),
            height=Inches(height_inches),
        )
    finally:
        if os.path.exists(fitted_path):
            os.remove(fitted_path)


def set_picture_in_front_of_text(inline):
    anchor = OxmlElement("wp:anchor")
    for attr, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251659264",
        "behindDoc": "0",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(attr, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")
    anchor.append(simple_pos)

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "column")
    align_h = OxmlElement("wp:align")
    align_h.text = "center"
    position_h.append(align_h)
    anchor.append(position_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "paragraph")
    offset_v = OxmlElement("wp:posOffset")
    offset_v.text = "0"
    position_v.append(offset_v)
    anchor.append(position_v)

    for tag in ("wp:extent", "wp:effectExtent"):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(element)

    wrap_none = OxmlElement("wp:wrapNone")
    anchor.append(wrap_none)

    for tag in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
        element = inline.find(qn(tag))
        if element is not None:
            anchor.append(element)

    inline.getparent().replace(inline, anchor)


def add_photo_to_cell(cell, image_file):
    if image_file is None:
        return

    image_file.seek(0)
    with Image.open(image_file) as source:
        image = ImageOps.exif_transpose(source).convert("RGB")

    max_width_inches = 6.25
    max_height_inches = 5.25
    width_ratio = max_width_inches / image.width
    height_ratio = max_height_inches / image.height
    scale = min(width_ratio, height_ratio)
    width_inches = image.width * scale
    height_inches = image.height * scale

    fd, image_path = tempfile.mkstemp(suffix=".jpg")
    os.close(fd)
    try:
        image.save(image_path, quality=95)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ""

        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline_shape = paragraph.add_run().add_picture(
            image_path,
            width=Inches(width_inches),
            height=Inches(height_inches),
        )
        set_picture_in_front_of_text(inline_shape._inline)
    finally:
        if os.path.exists(image_path):
            os.remove(image_path)


def validate_template(doc):
    if len(doc.tables) < 8:
        return "O template_reprovados.docx precisa possuir as tabelas esperadas da máscara FOR-BPC-5112."
    return None


equipment_registry = load_equipment_registry()
equipment_options = sorted(equipment_registry)
client_registry = load_client_registry()
client_options = sorted(client_registry)

st.subheader("1. Identificação e Local")
col1, col2, col3 = st.columns(3)
cliente = col1.selectbox("Cliente", ["Selecione", *client_options, "Outro"], key="reprovado_cliente_item")
if cliente == "Selecione":
    cliente = ""
if cliente == "Outro":
    cliente = col1.text_input("Especifique o Cliente", "", key="reprovado_cliente_outro", placeholder="DOF SUBSEA BRASIL SERVIÇOS LTDA")

vessel_options = sorted(client_registry.get(cliente, {}))
embarcacao = col2.selectbox("Embarcação (Local do Teste)", ["Selecione", *vessel_options, "Outro"], key="reprovado_embarcacao_item")
if embarcacao == "Selecione":
    embarcacao = ""
if embarcacao == "Outro":
    embarcacao = col2.text_input("Especifique a Embarcação", "", key="reprovado_embarcacao_outro", placeholder="SKANDI CHIEFTAIN")

if st.session_state.get("reprovado_ultima_embarcacao_cliente") != (cliente, embarcacao):
    st.session_state["reprovado_endereco_cliente"] = client_registry.get(cliente, {}).get(embarcacao, "")
    st.session_state["reprovado_ultima_embarcacao_cliente"] = (cliente, embarcacao)
endereco = col3.text_input("Endereço", key="reprovado_endereco_cliente", placeholder="PORTO DO AÇU")

col4, col5 = st.columns(2)
numero_certificado = col4.text_input("Número do relatório", "", placeholder="68430204")
item_certificado = col5.text_input("Item(s)", "", placeholder="32")

st.subheader("2. Dados da Inspeção e Equipamento")
col7, col9 = st.columns(2)
data_inspecao = col7.date_input("Data da Inspeção", value=None, format="DD/MM/YYYY")
ns = col9.text_input("NS (Número de Série)", "", placeholder="AST220S/03")

col10, col11, col12 = st.columns(3)
item = col10.selectbox("Equipamento", ["Selecione", *equipment_options, "Outro"], key="reprovado_equipamento_item")
if item == "Selecione":
    item = ""
if item == "Outro":
    item = col10.text_input("Especifique o Equipamento", "", key="reprovado_equipamento_outro", placeholder="Manilha")

if st.session_state.get("reprovado_ultimo_equipamento_item") != item:
    equipment_config = equipment_registry.get(item, {})
    st.session_state["reprovado_norma_referencia"] = equipment_config.get("criterio_aceitacao", "")
    st.session_state["reprovado_materia_prima"] = equipment_config.get("materia_prima", "")
    st.session_state["reprovado_ultimo_equipamento_item"] = item

norma_referencia = col11.text_input("Norma de Referência", key="reprovado_norma_referencia", placeholder="ABNT NBR 15637-1")
materia_prima = col12.text_input("Matéria Prima", key="reprovado_materia_prima", placeholder="POLIESTER")

col13, col14, col15, col16 = st.columns(4)
capac = col13.number_input("Capacidade (Capac)", min_value=0.1, value=None, step=0.1, placeholder="3.00")
unidade_capacidade = col14.text_input("Unidade da capacidade", "", placeholder="TON")
dimensao = col15.text_input("Dimensão", "", placeholder="2 MTS")
quantidade = col16.number_input("Quantidade", min_value=1, value=None, step=1, placeholder="1")

descricao_insuficiencia = st.text_area(
    "Descrição da Insuficiência",
    "",
    placeholder="FOI REALIZADO A INSPEÇÃO NO EQUIPAMENTO E O MESMO FOI REPROVADO POR APRESENTAR...",
)
foto_reprovado = st.file_uploader(
    "Imagem para a página 2",
    type=["jpg", "jpeg", "png"],
)

quality_signature_files = list_signature_files(QUALITY_SIGNATURE_DIR)
assinatura_qualidade = st.selectbox(
    "Assinatura do Controle de Qualidade",
    [None, *quality_signature_files],
    format_func=signature_label,
    index=0,
) if quality_signature_files else None
if not quality_signature_files:
    st.warning(f"Nenhuma assinatura encontrada na pasta {QUALITY_SIGNATURE_DIR}.")

technician_signature_files = list_signature_files(TECHNICIAN_SIGNATURE_DIR)
assinatura_tecnico = st.selectbox(
    "Assinatura do Técnico de Inspeção",
    [None, *technician_signature_files],
    format_func=signature_label,
    index=0,
) if technician_signature_files else None
if not technician_signature_files:
    st.warning(f"Nenhuma assinatura encontrada na pasta {TECHNICIAN_SIGNATURE_DIR}.")

submit_button = st.button("Gerar Relatórios")

if submit_button:
    with st.spinner("Processando documento..."):
        numero_certificado = numero_certificado.strip().strip("-")
        itens_certificado = parse_certificate_items(item_certificado)
        unidade_capacidade = unidade_capacidade.strip().upper()
        norma_referencia = norma_referencia.strip()

        if not numero_certificado or not itens_certificado:
            st.error("Preencha o Número do relatório e pelo menos um Item para gerar o relatório.")
            st.stop()
        if data_inspecao is None:
            st.error("Preencha a Data da Inspeção para gerar o relatório.")
            st.stop()
        if not cliente or not embarcacao or not endereco:
            st.error("Preencha Cliente, Embarcação e Endereço para gerar o relatório.")
            st.stop()
        if not item:
            st.error("Preencha o Equipamento para gerar o relatório.")
            st.stop()
        if not norma_referencia:
            st.error("Preencha a Norma de Referência para gerar o relatório.")
            st.stop()
        if not materia_prima:
            st.error("Preencha a Matéria Prima para gerar o relatório.")
            st.stop()
        if capac is None or not unidade_capacidade or not dimensao or quantidade is None:
            st.error("Preencha Capacidade, Unidade, Dimensão e Quantidade para gerar o relatório.")
            st.stop()

        data_str = data_inspecao.strftime("%d/%m/%Y")
        data_arquivo_str = data_inspecao.strftime("%d-%m-%Y")
        equipamento = item
        if unidade_capacidade == "TON":
            carga_trabalho = f"{int(capac * 1000):,} KG".replace(",", ".")
        else:
            carga_trabalho = f"{capac:g} {unidade_capacidade}"

        template_path = os.path.join(APP_DIR, "template_reprovados.docx")
        if not os.path.exists(template_path):
            st.error(f"Erro: Modelo 'template_reprovados.docx' não encontrado na pasta {APP_DIR}.")
            st.stop()

        with tempfile.TemporaryDirectory() as temp_dir:
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                for item_certificado_atual in itens_certificado:
                    certif = f"{numero_certificado}-{item_certificado_atual}-RI"
                    nome_base_arquivo = " ".join(
                        safe_filename_part(part)
                        for part in (certif, "REPROVADO", item, data_arquivo_str)
                    )

                    try:
                        doc = docx.Document(template_path)
                    except Exception as error:
                        st.error(f"Erro ao abrir o template_reprovados.docx: {error}")
                        st.stop()

                    template_error = validate_template(doc)
                    if template_error:
                        st.error(template_error)
                        st.stop()

                    replacements = {
                        "CONTROLE DE QUALIDADE (Quality Control):": "",
                        "RELATÓRIO Nº (Number Report):": certif,
                        "CLIENTE (Client):": cliente,
                        "CONTRATO (Contract):": "N/A",
                        "LOCAL DO TESTE (Local of the Test):": embarcacao,
                        "ENDEREÇO (Address):": endereco,
                        "EQUIPAMENTO (Equipment):": equipamento,
                        "SÉRIE (Serial Number):": ns,
                        "DATA DA INSPEÇÃO (Date):": data_str,
                        "NOTA FISCAL (Invoice):": "",
                        "NORMA DE REFERÊNCIA (Normative Reference):": norma_referencia,
                        "PROCEDIMENTO (Procedure):": "PRO-BPC-5.1 Rev.: 11",
                    }

                    for paragraph in doc.paragraphs:
                        replace_in_paragraphs([paragraph], replacements, data_str)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                replace_in_paragraphs(cell.paragraphs, replacements, data_str)

                    table_details = doc.tables[1]
                    set_cell_text(table_details.rows[1].cells[1], dimensao)
                    set_cell_text(table_details.rows[2].cells[1], materia_prima)
                    set_cell_text(table_details.rows[3].cells[1], carga_trabalho)
                    set_cell_text(table_details.rows[4].cells[1], f"{quantidade:02d} UNIDADE")

                    if descricao_insuficiencia.strip():
                        set_cell_text(doc.tables[2].rows[1].cells[0], descricao_insuficiencia)

                    add_photo_to_cell(doc.tables[5].rows[0].cells[0], foto_reprovado)

                    for table_index in (4, 7):
                        signature_table = doc.tables[table_index]
                        set_cell_text(signature_table.rows[0].cells[1], data_str)
                        set_cell_text(signature_table.rows[0].cells[3], data_str)
                        add_signature_to_cell(signature_table.rows[0].cells[0], assinatura_qualidade)
                        add_signature_to_cell(signature_table.rows[0].cells[4], assinatura_tecnico)

                    remove_empty_trailing_paragraphs(doc)

                    output_docx = os.path.join(temp_dir, f"{nome_base_arquivo}.docx")
                    output_pdf = os.path.join(temp_dir, f"{nome_base_arquivo}.pdf")
                    doc.save(output_docx)

                    try:
                        convert_to_pdf(output_docx, output_pdf)
                    except Exception as error:
                        st.error(format_conversion_error(error))
                        st.stop()

                    zip_file.write(output_docx, os.path.basename(output_docx))
                    zip_file.write(output_pdf, os.path.basename(output_pdf))

            zip_buffer.seek(0)

        st.success(f"{len(itens_certificado)} relatório(s) de reprovados gerado(s) com sucesso!")
        st.download_button(
            label="Baixar Word e PDF",
            data=zip_buffer.getvalue(),
            file_name=f"Relatorios_Reprovados_{numero_certificado}.zip",
            mime="application/zip",
        )
